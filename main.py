import csv
from datetime import date
import re

import pandas as pd
import streamlit as sl
from docx import Document
from scipy import stats as st

import specimen_bank
import qcmain


class TestClass:
    '''
    Class used to store all test data relevant to each specimen tested, 
    including replicate values and all stats. Replicates are grouped under
    the shared specimen used. For example, if specimen Boca124 is tested
    with three replicates, there will be one object for this specimen
    containing all three replicates saved, as well as stats such as mean
    and standard deviation of the replicates. Strip scoring data such as
    search and scoring window coordinates are also stored. 
    '''
    def __init__(self, formatted_data, rep_qty, index, c=0, condition=""):  

        self.fd = formatted_data
        self.qty = rep_qty
        self.index = index
        self.ltr_pos = []
        self.ver_pos = []

        # add line scoring positions
        for i in range(3):
            self.ltr_pos.append(formatted_data[index+i][5])
            self.ver_pos.append(formatted_data[index+i][4])

        if reps_in_title == True:
            self.name = self.fd[self.index][0][c:-2]  
        else: 
            self.fd[self.index][0][c:]
        self.condition = condition

        # Control line stats
        ctrl_data = []
        
        for j in range(self.qty):
            ctrl_data.append(round(float(self.fd[self.index + j][1]), 3))

        self.ctrl_mean = round(stats.mean(ctrl_data), 3)
        self.ctrl_median = round(stats.median(ctrl_data), 3)
        self.ctrl_SD = round(stats.stdev(ctrl_data), 3)
        self.ctrl_CV = round((self.ctrl_SD / self.ctrl_mean) * 100, 1)
        ctrl_MAD_list = [abs(k - self.ctrl_median) for k in ctrl_data]
        self.ctrl_MAD = round(stats.median(ctrl_MAD_list), 3)
        self.ctrl_data = ctrl_data
        
        # VER line stats
        ver_data = []
        
        try:
            for j in range(self.qty):
                ver_data.append(round(float(self.fd[self.index + j][2]), 3))
        except IndexError:
            sl.error(
                    '''Could not format data. A replicate may be duplicated 
                        or missing for {}.'''.format(self.name)
                        )
            
        self.ver_mean = round(stats.mean(ver_data), 3)
        self.ver_median = round(stats.median(ver_data), 3)
        self.ver_SD = round(stats.stdev(ver_data), 3)
        self.ver_CV = round((self.ver_SD / self.ver_mean) * 100, 1)
        ver_MAD_list = [abs(k - self.ver_median) for k in ver_data]
        self.ver_MAD = round(stats.median(ver_MAD_list), 3)
        self.ver_data = ver_data

        # LT/R line stats
        ltr_data = []
        for j in range(self.qty):
            ltr_data.append(round(float(self.fd[self.index + j][3]), 3))

        self.ltr_mean = round(stats.mean(ltr_data), 3)
        self.ltr_median = round(stats.median(ltr_data), 3)
        self.ltr_SD = round(stats.stdev(ltr_data), 3)
        self.ltr_CV = round((self.ltr_SD / self.ltr_mean) * 100, 1)
        ltr_MAD_list = [abs(k - self.ltr_median) for k in ltr_data]
        self.ltr_MAD = round(stats.median(ltr_MAD_list), 3)
        self.ltr_data = ltr_data


def get_conditions(read_in: list) -> list:
    '''
    Parse test ID's from imported csv file for experimental conditions.
    Conditions must follow naming convention in documentation.
    '''
    sample_ID = read_in['Test ID'].values
    test_type = read_in['Test Type'].values
    
    pattern = '[A-Za-z0-9]+ '
    condition_list = []
    
    for i, test in enumerate(sample_ID):
        if test_type[i] == 'SediaBio HIV - Rapid Recency - v8':
            result = re.findall(pattern, test)
            if len(result) > 0:
                find = result[0]
                condition = find.replace(' ', '')
                if condition not in condition_list:
                    condition_list.append(condition)
            
    return condition_list
        

def shapiro(data: list) -> str:
    '''
    Perform a Shapiro-Wilk test on test line values to test if values
    come from a normal distribution. Return the p-value. 
    '''
    result = st.shapiro(data)
    return result[1] # p value


def remove_nan(data_matrix: list) -> list:
    '''
    Removes any NA values from the data. 
    '''
    dm = []
    for i in data_matrix:
        if pd.isna(i[1]):
            pass
        else:
            dm.append(i)
    return dm


def remove_irrelevant_testing(data_matrix: list, panel: list, condition_list: list) -> list:
    '''
    Uses the panel tested and list of experimental conditions to remove irrelevant
    data from the data matrix. This was implemented because data from multiple 
    users can be pooled in the exported data by the test strip reader. 

    This function can be bypassed if desired by setting the "is there other data
    present" to no in the network GUI. 
    '''
    relevant_data = []
    if multiple_conditions == True:
        # remove tests that are not in panel.
        if reps_in_title == True:
            for condition in condition_list:
                c = len(condition) + 1
                for test in data_matrix:
                    if test[0][c:-2] in panel and condition == test[0][:c - 1]:
                        relevant_data.append(test)            
            return relevant_data
        else:
            for condition in condition_list:
                c = len(condition) + 1
                for test in data_matrix:
                    if test[0][c:] in panel and condition in test[0]:
                        relevant_data.append(test)
            return relevant_data
        
    else:
        if reps_in_title == True:
            for test in data_matrix:
                if test[0][:-2] in panel:
                    relevant_data.append(test)

            return relevant_data
        else:
            for test in data_matrix:
                if test[0] in panel:
                    relevant_data.append(test)
            return relevant_data


def format_data(
    dm: list, rep_qty: int, lots: list, condition_list: list, lines: list, rep_stats: list
    ) -> list:
    '''
    Creates test objects as an instance of the TestClass for each specimen run.
    Returns a formatted table for verification and LTR values as atributes of 
    the test object.
    '''
    test_objects = []

    data_for_csv = []
    headers = []
    cc_message = []
    ct_message = []

    if multiple_conditions == True:

        for line in lines:
            headers.append('Condition')
            headers.append(" ".join([line, '|', 'Sample']))
            for lot in lots:
                headers.append(lot)
            for stat in rep_stats:
                headers.append(stat)
            headers.append(' ')

        for condition in condition_list:
            c = len(condition) + 1

            for i in range(0, len(dm) - 2, rep_qty):
                if condition == dm[i][0][:c - 1]:
                    
                    test = TestClass(dm, rep_qty, i, c, condition)

                    test_objects.append(test)
                    data = []

                # CTRL vals                  
                    data.append(test.condition)
                    data.append(test.name)
                    for d in test.ctrl_data:
                        data.append(d)

                    ctrl_stats = {
                        'Mean':test.ctrl_mean, 'SD':test.ctrl_SD, 
                        '%CV':test.ctrl_CV, 'Median':test.ctrl_median, 
                        'MAD':test.ctrl_MAD
                        }
                    for stat in rep_stats:
                        data.append(ctrl_stats[stat])

                    data.append(' ')

                    # VER values
                    data.append(test.condition)
                    data.append(test.name)
                    for d in test.ver_data:
                        data.append(d)

                    ver_stats = {
                        'Mean':test.ver_mean, 'SD':test.ver_SD, 
                        '%CV':test.ver_CV, 'Median':test.ver_median, 
                        'MAD':test.ver_MAD
                        }
                    for stat in rep_stats:
                        data.append(ver_stats[stat])

                    data.append(' ')

                    # LTR values
                    if test_lines == 2:
                        data.append(test.condition)
                        data.append(test.name)
                        for n in test.ltr_data:
                            data.append(n)

                        ltr_stats = {
                            'Mean':test.ltr_mean, 'SD':test.ltr_SD, 
                            '%CV':test.ltr_CV, 'Median':test.ltr_median, 
                            'MAD':test.ltr_MAD
                            }
                        for stat in rep_stats:
                            data.append(ltr_stats[stat])
                        

                if len(data) > 0 and data not in data_for_csv:
                    data_for_csv.append(data)
                    
            data_for_csv.append([' '])

    elif multiple_conditions == False:
        for line in lines:
            headers.append(" ".join([line, '|', 'Sample']))
            for lot in lots:
                headers.append(lot)
            for stat in rep_stats:
                headers.append(stat)
            headers.append(' ')

        # test_objects = []
        condition = ""
        c = 0

        for i in range(0, len(dm) - 2, rep_qty):

            test = TestClass(dm, rep_qty, i, c, condition)
            test_objects.append(test)

            data = []

            #ctrl values
            data.append(test.name)
            for n in test.ctrl_data:
                data.append(n)
            ctrl_stats = {
                'Mean':test.ctrl_mean, 'SD':test.ctrl_SD, '%CV':test.ctrl_CV, 
                'Median':test.ctrl_median, 'MAD':test.ctrl_MAD
                }
            for stat in rep_stats:
                data.append(ctrl_stats[stat])
            data.append(' ')
            
            
            #Ver Values
            data.append(test.name)
            for n in test.ver_data:
                data.append(n)
            ver_stats = {
                'Mean':test.ver_mean, 'SD':test.ver_SD, '%CV':test.ver_CV, 
                'Median':test.ver_median, 'MAD':test.ver_MAD
                }
            for stat in rep_stats:
                data.append(ver_stats[stat])
            data.append(' ')

            # LTR values
            if test_lines == 2:
                data.append(test.name)
                for n in test.ltr_data:
                    data.append(n)
                ltr_stats = {
                    'Mean':test.ltr_mean, 'SD':test.ltr_SD, '%CV':test.ltr_CV,
                    'Median':test.ltr_median, 'MAD':test.ltr_MAD
                    }
                for stat in rep_stats:
                    data.append(ltr_stats[stat])

            if len(data) > 0:
                data_for_csv.append(data)

        data_for_csv.append([' '])


    # Check for color creep
    for test in test_objects:
        for i,_ in enumerate(test.ltr_pos):
            if int(test.ltr_pos[i]) > 530:
                cc_message.append(test.condition + ' ' + test.name + ' strip #' + str(i+1))
        for i,_ in enumerate(test.ver_pos):
            if int(test.ver_pos[i]) > 330:
                ct_message.append(test.condition + ' ' + test.name + ' strip #' + str(i+1))
            
    if len(cc_message) > 0:
        sl.error(
            '''Significant position shift for LTR line detected in the 
            following samples: {}. Color creep or covertape may have been 
            mistaken for the LTR line.'''.format(', '.join(cc_message)))
    if len(ct_message) > 0:
        sl.error(
            '''Significant position shift for VER line detected in the 
            following samples: {}. Covertape is likely in strip image.
            '''.format(', '.join(ct_message)))

    return headers, data_for_csv, test_objects


# generate test strip image doc from date and time of test
def generate_tmf901b(
        df, filepath, panel, test_lines, rep_qty, condition_list, date_comp, 
        version
        ):


    clen = len(condition_list) if len(condition_list) > 0 else 1

    sl.spinner('Generating TMF-901B...')
    time = df['Time Acquired'].values
    date = df['Test Date'].values
    test_no = df['Test Number'].values
    test_ID = df['Test ID'].values
    ver = df['Decision Title 2'].values
    ver_value = df['Decision Message 2'].values
    expected_strip_ct = len(panel) * rep_qty * clen

    if test_lines == 2:
        ltr = df['Decision Title 3'].values
        ltr_value = df['Decision Message 3'].values
    test_type = df['Test Type'].values

    test_to_ignore = [
        'Instrument Check, Visual, Red/Gold', 'ROF QC Test - Negative', 
        'ROF QC Test - Positive', 'QC Test - Negative', 'QC Test - Recent', 
        'QC Test - Long Term'
        ]

    document = Document(
        filepath + '''TMF-901B-00, Worksheet, Asanté™ HIV Rapid Recency®, 
        Reader Results Images.docx''')
    
    #date stamp
    head_table = document.tables[0]
    head_table.cell(0,4).text = '''Completed with Data Processing Pipeline {} 
            on {}.".format(version, date_comp)'''
    
    tbl = document.tables[1]


    # progress bar
    latest_iteration = sl.empty()
    bar = sl.progress(0)
    pc = 100/100/(len(panel)*rep_qty*clen)
    pic_ct = 0
    prog_bar_val = 0

    skipped_tests = []
    for i, _ in enumerate(time):
        if test_type[i] in test_to_ignore:  
            continue
        else:
            # create filename
            d = str(date[i][2:])
            t = str(time[i])

            filename = d.replace('-', '') + '_' + t.replace(':', '_')

            # Test No.
            row_cells = tbl.add_row().cells
            paragraph = row_cells[0].paragraphs[0]
            run = paragraph.add_run(str(test_no[i]))

            # Test Date
            paragraph = row_cells[1].paragraphs[0]
            run = paragraph.add_run(date[i])

            # Test ID
            paragraph = row_cells[2].paragraphs[0]
            run = paragraph.add_run(test_ID[i])

            # Line name
            paragraph = row_cells[3].paragraphs[0]
            run = paragraph.add_run()
            run = paragraph.add_run(str(ver[i]))

            # line value
            paragraph = row_cells[4].paragraphs[0]
            run = paragraph.add_run()
            run = paragraph.add_run(str(round(ver_value[i], 3)))

            if test_lines == 2:
                # line 2 name
                paragraph = row_cells[5].paragraphs[0]
                run = paragraph.add_run()
                run = paragraph.add_run(str(ltr[i]))

                # line 2 value
                paragraph = row_cells[6].paragraphs[0]
                run = paragraph.add_run()
                run = paragraph.add_run(str(round(ltr_value[i], 3)))


            try:
                # strip image
                paragraph = row_cells[-3].paragraphs[0]
                run = paragraph.add_run()
                run.add_picture(
                    filepath + 'strip_images/' + filename + '/Strip.jpg',
                    width=3100000, height=660000
                    )
                a, b, c = row_cells[-3:]
                a.merge(b)
                a.merge(c)
                # Update the progress bar with each iteration.
                latest_iteration.text('Strip images moved: {} of {}'.format(
                    pic_ct, expected_strip_ct)
                    )
                bar.progress(prog_bar_val)
                prog_bar_val += pc
                pic_ct += 1
            except:
                skipped_tests.append(test_ID[i])

    bar.progress(100)

    if len(skipped_tests) > 0:
        sl.error(
            "Could not find strip images for the following tests: {}.".format(
                ', '.join(skipped_tests)))
        sl.error('''Verify that Test Date is in YYYY-MM-DD format and that 
                 Time Acquired is in HH:MM:SS 24hr format in data.csv.'''
                 )

    doc_path = filepath + '{} Completed TMF-910B.docx'.format(study_name)
    document.save(doc_path)
    sl.info('TMF-901B Exported.')


    if pic_ct < expected_strip_ct:
        sl.error(
            '''{} strip images were expected based on user parameters but only
            {} were identified. Check data.csv and strip_images folder for 
            errors.'''.format(expected_strip_ct, pic_ct)
            )

           
# main function calls
def main():
    
    # file read in.
    raw_data = read_in[[
        'Test ID', 'Decision Message 1', 'Decision Message 2', 
        'Decision Message 3', 'Position.1', 'Position.2'
        ]].values
    data_matrix = remove_nan(raw_data)


    # formatting parameters.
    lots_list = test_header.split(', ')
    
    if test_lines == 2:
        lines = ['CTRL', 'VER', 'LTR']
    elif test_lines == 1:
        lines = ['CTRL', 'VER']

    # function calls to handle data file.
    sl.spinner('Formatting Data...')
    if other_data:
        relevant_data = remove_irrelevant_testing(
            data_matrix, panel, condition_list
            )  

        
        headers, data_for_csv, test_objects = format_data(
            relevant_data, rep_qty, lots_list, condition_list, lines, rep_stats
            )
    else:
        headers, data_for_csv, test_objects = format_data(
            data_matrix, rep_qty, lots_list, condition_list, lines, rep_stats
            )
    sl.info('Data formatted.')
    

    

    # writing to csv file
    if 'Export Formatted CSV file' in docs_to_export:
        sl.spinner('Exporting data...')
        output = filepath + '{}_Data_analysis.csv'.format(study_name)
        with open(output, 'w', newline='') as csvfile:
            csvwriter = csv.writer(csvfile)
            csvwriter.writerow(headers)
            csvwriter.writerows(data_for_csv)
            sl.success('Exported formatted CSV.')
            
    if 'Download formatted data as txt' in docs_to_export:
        csv_data = ''
        csv_data += '\t'.join(headers)
        csv_data += '\n'
        for i in data_for_csv:
            csv_data += '\t'.join([str(j) for j in i])
            csv_data += '\n'
        
        sl.download_button(
            'Download formatted data', csv_data, 
            '{} Formatted data'.format(study_name))


    # Generate TMF-901B Doc.
    if 'TMF-901B' in docs_to_export:
        generate_tmf901b(
            read_in, filepath, panel, test_lines, rep_qty, 
            condition_list, date_comp, version
            )


    if "Balloons" in docs_to_export:
        sl.balloons()


version = "v1.4"



# User parameters and streamlit app appearance
sl.sidebar.write("""
# Sedia Data Processing Pipeline {}


""".format(version))
    
pipeline = sl.sidebar.radio('Select pipeline: ', ['PDS', 'QC'])

if pipeline == 'PDS':
    
    # User configurations in web app.
    sl.write("""
    # PDS Data Processing
    
    Configure the following parameters:
    
    """)

    test = sl.radio("Select an assay", ['Rapid Recency', 'Oral Fluid'])
    test_lines = 2 if test == 'Rapid Recency' else 1
    
    docs_to_export = sl.multiselect(
        'Select documents to generate:', 
        ['Download formatted data as txt', 'Export Formatted CSV file',
         'TMF-901B',  'Balloons']
        )
    
    if 'Export Formatted CSV file' in docs_to_export:
        filepath = sl.text_input("Filepath: ", "")
        filepath += "/"
        read_in = pd.read_csv(f"{filepath}data.csv", encoding="cp1252", engine="python")
        
    if 'Download formatted data as txt' in docs_to_export:
        data_readin = sl.file_uploader("Upload TestResults.csv")
        if data_readin != None:
            read_in = pd.read_csv(data_readin)
    
    study_name = sl.text_input("Study name/number: ", "")
    
    rep_qty = sl.number_input("Number of technical replicates:", min_value=1, max_value=9, step=1)
    
    reps_input = sl.radio(
        "Are replicate numbers listed in the Test ID?", ['Yes', 'No']
        )
    reps_in_title = True if reps_input == 'Yes' else False
    default_header = 'rep 1, rep 2, rep 3'
    
    stats_input = sl.multiselect(
        'Select statistics to run on replicates:', ['Mean', 'SD', '%CV', 'Median', 'MAD']
        )
    
    stat_key = ['Mean', 'SD', '%CV', 'Median', 'MAD']
    order = {v:i for i,v in enumerate(stat_key)}
    rep_stats = sorted(stats_input, key=lambda x: order[x])
    stats_len = len(rep_stats)
    
    
    test_header = sl.text_input("Header for data table: ", default_header)
    
    
    # Experimental condition selection
    multiple_conditions_input = sl.radio(
        'Experimental conditions', ['I don\'t have experimental conditions', 'Search for my conditions', 'Enter conditions']
        )
    
    if multiple_conditions_input == 'Search for my conditions':
        multiple_conditions = True
        condition_list = get_conditions(read_in)
        conditions = sl.multiselect('Select your conditions: ', condition_list)
    elif multiple_conditions_input == 'Enter conditions':
        multiple_conditions = True
        conditions = sl.text_input('Type conditions separated by comma+space:')
        condition_list = conditions.split(', ')
    elif multiple_conditions_input == 'No experimental conditions':
        multiple_conditions = False
    
    other_data_input = sl.radio(
        "Is other data present in data set?", ['Yes', 'No']
        )
    other_data = True if other_data_input == 'Yes' else False
    
    panel_select = sl.radio('', ['Select from supported panels', 'Input specimen'])
    if panel_select == 'Input specimen':
        specimens_tested = sl.text_input("Specimen tested or panel: ", "")
        panel = specimens_tested.split(', ')
    elif panel_select == 'Select from supported panels':
        specimens_tested = sl.selectbox('Supported panels:', ['9169', '9170', 
                                                              '9172', 'Prozone', 
                                                              'Gilead', 'Gilead 2', 
                                                              'RR', 'flex', '9172/9168 Hybrid'])
        panel = specimen_bank.get_specimens(specimens_tested)
    
    
    sl.write("Last updated 11 NOV 2021.")
    
    today = date.today()
    date_comp = today.strftime("%d-%b-%Y")
    
    done = sl.button('Done')
    
    
    if done == True:
        main()
        
elif pipeline == 'QC':
    qcmain.main()
