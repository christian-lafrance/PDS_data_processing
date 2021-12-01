import streamlit as sl
import pandas as pd
import re
from docx import Document
from pathlib import Path

# sheet_names = ['1. Overview', '2. Data Entry', '3. Lot Consistency & CV', 
#                '4. Lot Consistency & CV Plots', '5. Plots', '6. Specificity', 
#                '7. Expected Data', '8. Ref Historical Data', 'Raw Data Edited', 
#                'Raw Data Unedited', 'Raw Data visual']
# current_test_type_package = 'SediaBio HIV - Rapid Recency - v8'


class QcClass:
    
    def __init__(self, test_ID, test_num, test_type, ltr_pos, ctrl_val, 
                        ver_val, ltr_val, time, date):
    
        self.test_ID = test_ID
        print(test_ID)
        
        sublot_pattern = 'sublot_[A-Z].[A-Z] '
        #pm_pattern = '9169 [0-7][0-9]'
        pm_pattern = 'Sample [0-9]'
        neg_pat = 'N [0-9]+'
        
        
        if self.test_ID.startswith('N'):
            self.sample = re.findall(neg_pat, test_ID)[0]
            self.sublot = 'A' # re.findall(sublot_pattern, test_ID)[0]
        else:
            self.sublot = re.findall(sublot_pattern, test_ID)[0]
            self.sample = re.findall(pm_pattern, test_ID)[0]
        
        
        self.num = test_num
        self.test_type = test_type
        self.ltr_pos = ltr_pos
        self.ctrl_val = ctrl_val
        self.ver_val = ver_val
        self.ltr_val = ltr_val
        self.time = time
        self.date = date
        
    def generate_foldername(self):
        d = str(self.date[2:])
        t = str(self.time)
    
        folder_name = d.replace('-', '') + '_' + t.replace(':', '_')
        
        return folder_name
        


def image(foldernames, strip_images):
    
    strip_image_dict = {}
    
    for i, image in enumerate(strip_images):
        strip_image_dict[foldernames[i][-16:-1]] = image
    
        print(strip_image_dict)
    return strip_image_dict
    

def inst(
        test_ID, test_num, test_type, ltr_pos, ctrl_val, 
                        ver_val, ltr_val, strip_image_dict, time, date
                        ):        
    test_objects = []
    retests = []
    retest_pat = 'retest'
    
    for i, test in enumerate(test_ID):
        if 'buffer' not in test_ID[i]:
            
            temp_test = QcClass(test_ID[i], test_num[i], test_type[i], ltr_pos[i], ctrl_val[i], 
                            ver_val[i], ltr_val[i], time[i], date[i])
            
            if retest_pat in temp_test.test_ID:
                retests.append(temp_test)
            else:
                test_objects.append(temp_test)

    for rt in retests:
        for i, test in enumerate(test_objects):
            if rt.sample == test.sample and rt.test_ID != test.test_ID and rt.sublot == test.sublot:
                    test_objects[i] = rt

    return test_objects



# def excel_handle(test_objects, ss, qsf_file_name):
    
    
#     sublot_1 = []
#     sublot_2 = []
#     sublot_3 = []
    
    
#     xfile = openpyxl.load_workbook(ss)

#     sheet = xfile.get_sheet_by_name(sheet_names[1])
    
#     j = 6 # starting location in spreadsheet for N
#     k = 6 # starting location in spreadsheet for 9169
#     for i, test in enumerate(test_objects):
#         if 'N' in test.sample:
#             sheet[f'H{j}'] = test.ver_val
#             sheet[f'I{j}'] = test.ltr_val
#             j += 1
#         elif '9169' in test.sample:
#             sheet[f'C{k}'] = test.ver_val
#             sheet[f'D{k}'] = test.ltr_val
#             k += 3
#         if k == 222:
#             k = 7
#         elif k == 223:
#             k = 8

#     xfile.save("{}.xlsx".format(qsf_file_name))

    
def read_in(csv, file_name, foldernames, strip_images):
    test_ID = csv['Test ID'].values
    test_num = csv['Test No'].values
    test_type = csv['Test Type'].values
    ltr_pos = csv['Position.2'].values
    ctrl_val = csv['Decision Message 1'].values
    ver_val = csv['Decision Message 2'].values
    ltr_val = csv['Decision Message 3'].values
    time = csv['Time Acquired'].values
    date = csv['Test Date'].values
    
    # returns dictionary with foldername as key and image as value. 
    strip_image_dict = image(foldernames, strip_images)
    
    test_objects = inst(
                        test_ID, test_num, test_type, ltr_pos, ctrl_val, 
                        ver_val, ltr_val, strip_image_dict, time, date
                        )
    
    # excel_handle(test_objects, ss, qsf_file_name)
    
    return test_objects, strip_image_dict
    
        

def generate_tmf901b(test_objects, tmf901b, strip_image_dict, filepath):

    document = Document(tmf901b)
    
    tbl = document.tables[1]

    # progress bar
    latest_iteration = sl.empty()
    bar = sl.progress(0)
    pc = 100/100/(len(test_objects))
    expected_strip_ct = len(test_objects)
    pic_ct = 1
    prog_bar_val = 0

    skipped_tests = []
    for i, test in enumerate(test_objects):


        # Test No.
        row_cells = tbl.add_row().cells
        paragraph = row_cells[0].paragraphs[0]
        run = paragraph.add_run(str(test.num))

        # Test Date
        paragraph = row_cells[1].paragraphs[0]
        run = paragraph.add_run(test.date)

        # Test ID
        paragraph = row_cells[2].paragraphs[0]
        run = paragraph.add_run(test.test_ID)

        # Line name
        paragraph = row_cells[3].paragraphs[0]
        run = paragraph.add_run()
        run = paragraph.add_run('Verification Line')

        # line value
        paragraph = row_cells[4].paragraphs[0]
        run = paragraph.add_run()
        run = paragraph.add_run(str(round(test.ver_val, 3)))
        
        # Line name
        paragraph = row_cells[5].paragraphs[0]
        run = paragraph.add_run()
        run = paragraph.add_run('LT/R Line')

        # line value
        paragraph = row_cells[6].paragraphs[0]
        run = paragraph.add_run()
        run = paragraph.add_run(str(round(test.ltr_val, 3)))


        # try:
        # strip image
        paragraph = row_cells[-3].paragraphs[0]
        run = paragraph.add_run()
        #run.add_picture(test.strip_image, width=3100000, height=660000)
        root = Path(__file__).parent
        picture = os.path.join(root, (filepath + '/' + test.generate_foldername() + '/Strip.jpg')
        run.add_picture(picture, width=3100000, height=660000)
        # 'Streamlit_rakuten/' + 
        
        a, b, c = row_cells[-3:]
        a.merge(b)
        a.merge(c)
        # Update the progress bar with each iteration.
        latest_iteration.text('Strip images moved: {} of {}'.format(
            pic_ct, expected_strip_ct)
            )
        bar.progress(prog_bar_val)
        prog_bar_val += pc
        pic_ct += 1
        # except:
        #     skipped_tests.append(test.test_ID)

    bar.progress(100)

    if len(skipped_tests) > 0:
        sl.error(
            "Could not find strip images for the following tests: {}.".format(
                ', '.join(skipped_tests)))
        sl.error('''Verify that Test Date is in YYYY-MM-DD format and that 
                 Time Acquired is in HH:MM:SS 24hr format in data.csv.'''
                 )

    #document.save('test file.docx')
    #sl.download_button('Download TMF-901B', document, file_name='TMF-90B.txt')
             
    return document




# streamlit interface
def main():
    
    # User configurations in web app.
    sl.write("""
    # QC Data Formatting
    
    Configure the following parameters:
    
    """)
    
    # assay selection
    test = sl.radio("Select an assay", ['Rapid Recency', 'Oral Fluid'])
    test_lines = 2 if test == 'Rapid Recency' else 1
    
    # Read in data
    csv = sl.file_uploader("Upload TestResults.csv")
    
    tmf901b = sl.file_uploader('Upload blank TMF-901B:')
        
    strip_images = sl.file_uploader('Drag and drop all strip image folders together: ', accept_multiple_files=True)
    filepath = sl.text_input('Filepath to strip_images folder: ')
    
    
    foldernames = sl.text_input("Strip image folder names (Must be from USB drive): ").split(' ')
    
    file_name = "Completed file"
    
    
    balloons = sl.checkbox("Balloons")
    
    done = sl.button('Done')
    
    if done == True: 
        test_objects, strip_image_dict = read_in(
            pd.read_csv(csv), file_name, foldernames, strip_images
            )
        
        sl.text(strip_image_dict)
        
        file = generate_tmf901b(test_objects, tmf901b, strip_image_dict, filepath)
    
        file.save('test.docx')
        
        f = open('test.docx', 'rb')
        sl.download_button("Download completed TMF-901B", f)
        f.close()
        
        if balloons == True:
            sl.balloons()

        


